from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from template import DailyReport


def set_cell_padding(cell, top, bottom, left, right):
    """设置单元格内边距（单位cm）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(int(val * 360)))  # 1cm = 360 twips
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_cell_vertical_center(cell):
    """设置单元格垂直居中"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)


def build_docx(report: DailyReport, output_path: str) -> str:
    """将日报数据渲染成Word文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)

    # 标题
    title = doc.add_heading(f"{report.project_name} 施工日报", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"日期：{report.date}").font.size = Pt(12)

    doc.add_paragraph()  # 空行

    # 表格形式展示各字段
    fields = [
        ("天气情况", report.weather),
        ("施工部位", report.construction_location),
        ("主要工作内容", report.main_work),
        ("人员情况", report.personnel),
        ("机械设备", report.machinery),
        ("安全文明施工", report.safety),
        ("存在问题及处理措施", report.issues),
        ("明日工作计划", report.tomorrow_plan),
    ]

    table = doc.add_table(rows=len(fields), cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(fields):
        # 左列：字段名
        cell_label = table.cell(i, 0)
        cell_label.text = label
        cell_label.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in cell_label.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(11)
        cell_label.width = Cm(4)
        set_cell_padding(cell_label, 0.1, 0.1, 0.2, 0.2)
        set_cell_vertical_center(cell_label)

        # 右列：内容
        cell_value = table.cell(i, 1)
        cell_value.text = value
        cell_value.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in cell_value.paragraphs[0].runs:
            run.font.size = Pt(11)
        cell_value.width = Cm(12)
        set_cell_padding(cell_value, 0.1, 0.1, 0.35, 0.35)  # 左右各增加0.15cm
        set_cell_vertical_center(cell_value)

    doc.add_paragraph()  # 空行

    # 图片占位符
    if report.photo_placeholders:
        photo_heading = doc.add_heading("附图", level=2)
        for idx, desc in enumerate(report.photo_placeholders, 1):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f"【附图{idx}：{desc}】")
            run.font.size = Pt(11)
            run.font.color.rgb = None  # 默认颜色
            run.font.italic = True

    # 保存
    doc.save(output_path)
    return output_path
